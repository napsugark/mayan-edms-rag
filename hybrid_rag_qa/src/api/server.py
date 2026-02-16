"""
FastAPI server for Mayan EDMS integration

Exposes REST endpoints for:
- Document indexing via file upload (POST /index)
- AI-assisted search with permissions (POST /query)
- Health checks and status

Mayan EDMS pushes documents to /index with allowed_users.
RAG enforces permissions when querying via /query.

Run with: uvicorn src.api.server:app --reload
Or: poetry run api-server
"""

import io
import json
import hashlib
import logging
import tempfile
from typing import Optional, List, Dict
from datetime import datetime
from contextlib import asynccontextmanager
import asyncio
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, HTTPException, Depends, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field

from haystack import Document
from haystack.components.converters import PyPDFToDocument

# Import app components
from pathlib import Path

from ..app import HybridRAGApplication
from ..integrations.mayan_client import MayanClient
from .. import config

logger = logging.getLogger("HybridRAG.API")

# Supported file extensions for /index endpoint
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}

# =====================
# Global state
# =====================
rag_app: Optional[HybridRAGApplication] = None
executor = ThreadPoolExecutor(max_workers=3)

# =====================
# Pydantic models
# =====================

class MayanIndexResponse(BaseModel):
    """Response for POST /index endpoint (Mayan EDMS spec)"""
    status: str = Field(..., description="'indexed' on success")
    document_id: int
    document_version_id: int


class MayanQueryRequest(BaseModel):
    """Request body for POST /query endpoint (Mayan EDMS spec)"""
    user_id: int = Field(..., description="Requesting user's ID for permission filtering")
    query: str = Field(..., description="Search query text")


class MayanQueryResultItem(BaseModel):
    """Single result item in query response"""
    document_id: int
    document_version_id: int


class MayanQueryResponse(BaseModel):
    """Response for POST /query endpoint (Mayan EDMS spec)"""
    answer: str = Field(..., description="AI-generated answer")
    results: List[MayanQueryResultItem] = Field(
        default_factory=list, 
        description="List of matching document versions"
    )


class HealthResponse(BaseModel):
    """Health check response"""
    status: str
    timestamp: str
    rag_initialized: bool
    qdrant_connected: bool
    document_count: int


class StatusResponse(BaseModel):
    """Detailed status response"""
    qdrant_url: str
    collection: str
    document_count: int
    features: Dict[str, bool]


# =====================
# Lifespan management
# =====================

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize and cleanup resources"""
    global rag_app
    
    logger.info("=" * 60)
    logger.info("Starting Hybrid RAG API Server")
    logger.info("=" * 60)
    
    # Initialize RAG application
    try:
        rag_app = HybridRAGApplication()
        logger.info("RAG application initialized")
    except Exception as e:
        logger.error(f"Failed to initialize RAG application: {e}")
        raise
    
    yield
    
    # Cleanup
    logger.info("Shutting down API server")
    executor.shutdown(wait=True)


# =====================
# FastAPI app
# =====================

app = FastAPI(
    title="Hybrid RAG API",
    description="REST API for Mayan EDMS integration with Haystack RAG",
    version="1.0.0",
    lifespan=lifespan,
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =====================
# Dependencies
# =====================

def get_rag_app() -> HybridRAGApplication:
    """Dependency to get RAG application"""
    if rag_app is None:
        raise HTTPException(status_code=503, detail="RAG application not initialized")
    return rag_app


# =====================
# Helper functions
# =====================


def file_to_haystack_document(
    file_content: bytes,
    filename: str,
    document_id: int,
    document_version_id: int,
) -> Document:
    """
    Convert uploaded file bytes to a Haystack Document.
    
    Supports: PDF, DOCX, TXT, and image files.
    For images, we assume OCR has already been performed by Mayan.
    
    Args:
        file_content: Raw file bytes
        filename: Original filename (used for extension detection)
        document_id: Mayan document ID
        document_version_id: Mayan document version ID
        
    Returns:
        Haystack Document with content extracted from file
    """
    suffix = Path(filename).suffix.lower()
    
    if suffix == ".pdf":
        # Handle PDF using PyPDFToDocument
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_content)
            tmp_path = Path(tmp.name)
        
        try:
            converter = PyPDFToDocument()
            result = converter.run(sources=[tmp_path])
            documents = result.get("documents", [])
            
            if documents:
                # Combine all pages into one document
                combined_content = "\n\n".join(doc.content for doc in documents if doc.content)
                return Document(
                    content=combined_content,
                    meta={
                        "file_path": filename,
                        "source": f"mayan:{document_id}",
                    }
                )
            else:
                raise ValueError(f"Could not extract content from PDF: {filename}")
        finally:
            tmp_path.unlink(missing_ok=True)
    
    elif suffix == ".docx":
        # Handle DOCX - use python-docx if available
        try:
            import docx
            with io.BytesIO(file_content) as docx_stream:
                doc = docx.Document(docx_stream)
                content = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
                return Document(
                    content=content,
                    meta={
                        "file_path": filename,
                        "source": f"mayan:{document_id}",
                    }
                )
        except ImportError:
            raise HTTPException(
                status_code=422,
                detail="DOCX support requires python-docx. Install with: pip install python-docx"
            )
    
    elif suffix == ".txt":
        # Handle plain text
        try:
            content = file_content.decode("utf-8")
        except UnicodeDecodeError:
            content = file_content.decode("latin-1")
        
        return Document(
            content=content,
            meta={
                "file_path": filename,
                "source": f"mayan:{document_id}",
            }
        )
    
    elif suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        # Image file — fetch OCR content from Mayan EDMS API
        try:
            mayan_client = MayanClient(
                base_url=config.MAYAN_URL,
                api_token=config.MAYAN_API_TOKEN,
            )
            ocr_content = mayan_client.get_document_ocr_content(document_id)
            
            if not ocr_content or not ocr_content.strip():
                raise HTTPException(
                    status_code=422,
                    detail=f"No OCR content available for image {filename} "
                           f"(document_id={document_id}). "
                           f"Ensure Mayan has completed OCR processing."
                )
            
            return Document(
                content=ocr_content,
                meta={
                    "file_path": filename,
                    "source": f"mayan:{document_id}",
                    "content_source": "mayan_ocr",
                }
            )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=422,
                detail=f"Failed to fetch OCR content from Mayan for {filename}: {str(e)}"
            )
    
    else:
        raise HTTPException(
            status_code=422,
            detail=f"Unsupported file type: {suffix}. Supported: {SUPPORTED_EXTENSIONS}"
        )


# =====================
# Endpoints
# =====================

@app.get("/", response_model=Dict[str, str])
async def root():
    """Root endpoint"""
    return {
        "service": "Hybrid RAG API",
        "version": "1.0.0",
        "docs": "/docs",
    }


@app.get("/health", response_model=HealthResponse)
async def health_check(app: HybridRAGApplication = Depends(get_rag_app)):
    """Health check endpoint"""
    try:
        doc_count = app.get_document_count()
        qdrant_ok = True
    except Exception:
        doc_count = 0
        qdrant_ok = False
    
    return HealthResponse(
        status="healthy" if qdrant_ok else "degraded",
        timestamp=datetime.now().isoformat(),
        rag_initialized=rag_app is not None,
        qdrant_connected=qdrant_ok,
        document_count=doc_count,
    )


@app.get("/status", response_model=StatusResponse)
async def get_status(app: HybridRAGApplication = Depends(get_rag_app)):
    """Get detailed system status"""
    return StatusResponse(
        qdrant_url=app.qdrant_url,
        collection=app.collection_name,
        document_count=app.get_document_count(),
        features={
            "metadata_extraction": config.ENABLE_METADATA_EXTRACTION,
            "summarization": config.ENABLE_SUMMARIZATION,
            "reranking": config.USE_RERANKER,
        },
    )


@app.post("/index", response_model=MayanIndexResponse)
async def mayan_index_document(
    document_id: int = Form(..., description="Mayan document ID"),
    document_version_id: int = Form(..., description="Mayan document version ID"),
    allowed_users: str = Form(..., description="JSON array of user IDs with access"),
    file: UploadFile = File(..., description="Document file to index"),
    app: HybridRAGApplication = Depends(get_rag_app),
):
    """
    Index a document from Mayan EDMS via file upload.
    
    This endpoint accepts multipart/form-data with:
    - document_id: Mayan document ID
    - document_version_id: Mayan document version ID  
    - allowed_users: JSON array of user IDs (e.g., "[1, 5, 8]")
    - file: The document file (PDF, DOCX, TXT)
    
    Behavior:
    - Parses and validates all fields
    - If document_version_id exists, deletes existing vectors (idempotent overwrite)
    - Processes file through indexing pipeline (MetadataEnricher preserves existing logic)
    - Adds Mayan metadata (document_id, document_version_id, allowed_users)
    - Returns success response for Mayan EDMS integration
    """
    logger.info(
        f"POST /index: document_id={document_id}, "
        f"document_version_id={document_version_id}, file={file.filename}"
    )
    
    # Parse allowed_users from JSON string
    try:
        allowed_users_list = json.loads(allowed_users)
        if not isinstance(allowed_users_list, list):
            raise ValueError("allowed_users must be a JSON array")
        # Ensure all elements are integers
        allowed_users_list = [int(uid) for uid in allowed_users_list]
    except (json.JSONDecodeError, ValueError, TypeError) as e:
        logger.error(f"Invalid allowed_users format: {allowed_users}")
        raise HTTPException(
            status_code=422,
            detail=f"Invalid allowed_users format. Expected JSON array of integers. Error: {str(e)}"
        )
    
    # Validate file extension
    if not file.filename:
        raise HTTPException(status_code=422, detail="File must have a filename")
    
    suffix = Path(file.filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=422,
            detail=f"Unsupported file type: {suffix}. Supported: {list(SUPPORTED_EXTENSIONS)}"
        )
    
    try:
        # Read file content
        file_content = await file.read()
        
        if not file_content:
            raise HTTPException(status_code=422, detail="Empty file received")
        
        loop = asyncio.get_event_loop()
        
        # Convert file to Haystack document (runs in thread pool for I/O)
        def process_and_index():
            # Convert uploaded file to Haystack Document
            content_hash = hashlib.sha256(file_content).hexdigest()
            haystack_doc = file_to_haystack_document(
                file_content=file_content,
                filename=file.filename,
                document_id=document_id,
                document_version_id=document_version_id,
            )
            haystack_doc.meta["content_hash"] = content_hash
            
            # Index using the Mayan-specific method (skips if already indexed)
            result = app.index_mayan_document(
                document=haystack_doc,
                document_id=document_id,
                document_version_id=document_version_id,
                allowed_users=allowed_users_list,
            )
            
            return result
        
        result = await loop.run_in_executor(executor, process_and_index)
        
        return MayanIndexResponse(
            status=result["status"],  # "indexed" or "skipped"
            document_id=document_id,
            document_version_id=document_version_id,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to index document: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to index document: {str(e)}"
        )


@app.post("/query", response_model=MayanQueryResponse)
async def mayan_query(
    request: MayanQueryRequest,
    app: HybridRAGApplication = Depends(get_rag_app),
):
    """
    AI-assisted search endpoint for Mayan EDMS.
    
    This endpoint:
    1. Accepts JSON body with user_id and query
    2. Performs semantic search via Haystack/Qdrant
    3. Filters results by permission (user_id in allowed_users)
    4. Deduplicates by (document_id, document_version_id)
    5. Generates answer using only permitted documents
    6. Returns answer and list of matching document versions
    
    Security: Only documents where user_id is in allowed_users are returned.
    Mayan EDMS will re-filter, but RAG enforces permissions to prevent data leaks.
    """
    logger.info(f"POST /query: user_id={request.user_id}, query='{request.query}'")
    
    try:
        loop = asyncio.get_event_loop()
        
        # Run query with permission filtering in thread pool
        result = await loop.run_in_executor(
            executor,
            lambda: app.query_with_permissions(
                query=request.query,
                user_id=request.user_id,
            )
        )
        
        return MayanQueryResponse(
            answer=result.get("answer", ""),
            results=[
                MayanQueryResultItem(
                    document_id=item["document_id"],
                    document_version_id=item["document_version_id"],
                )
                for item in result.get("results", [])
            ],
        )
        
    except Exception as e:
        logger.error(f"Query failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Query failed: {str(e)}")



# =====================
# Server entry point
# =====================

def start_server(host: str = "0.0.0.0", port: int = 8000, reload: bool = False):
    """Start the FastAPI server"""
    import uvicorn
    uvicorn.run(
        "src.api.server:app",
        host=host,
        port=port,
        reload=reload,
        log_level="info",
    )


if __name__ == "__main__":
    start_server(reload=True)
